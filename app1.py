import streamlit as st
import pandas as pd
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Gọi thư viện kết nối Google Sheets của Streamlit
from streamlit_gsheets import GSheetsConnection

WORD_FILE = "tonghop.docx"

# Khởi tạo kết nối tới Google Sheets (Đã cấu hình trong mục Secrets)
conn = st.connection("gsheets", type=GSheetsConnection)

def luu_y_kien_vao_sheets(mang_bo_phan, ho_ten, chuc_vu, noi_dung):
    # 1. Đọc dữ liệu hiện tại đang có trên Google Sheets về dưới dạng DataFrame
    try:
        df_hien_tai = conn.read(ttl=0) # ttl=0 để luôn đọc dữ liệu mới nhất, không dùng bộ nhớ đệm
    except Exception:
        df_hien_tai = pd.DataFrame()

    # 2. Tạo bản ghi mới
    new_data = {
        "Thời gian": [datetime.now().strftime("%d/%m/%Y %H:%M:%S")],
        "Mảng/Bộ phận": [mang_bo_phan],
        "Họ và Tên": [ho_ten if ho_ten else "Ẩn danh"],
        "Chức vụ Đảng/Chính quyền": [chuc_vu if chuc_vu else "Đảng viên/Quần chúng"],
        "Nội dung ý kiến": [noi_dung]
    }
    df_new = pd.DataFrame(new_data)
    
    # 3. Gộp bản ghi mới vào dữ liệu cũ
    if not df_hien_tai.empty:
        # Loại bỏ các dòng trống hoàn toàn nếu có
        df_hien_tai = df_hien_tai.dropna(how='all')
        df_cap_nhat = pd.concat([df_hien_tai, df_new], ignore_index=True)
    else:
        df_cap_nhat = df_new

    # 4. Ghi đè DataFrame đã cập nhật ngược trở lại Google Sheets
    conn.update(data=df_cap_nhat)

def xuat_file_word_tu_sheets(df):
    if df.empty:
        return False
        
    doc = Document()
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("BÁO CÁO TỔNG HỢP Ý KIẾN ĐÓNG GÓP\n")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(13)
    run_title.bold = True
    
    run_subtitle = title_p.add_run(f"Phục vụ sinh hoạt Chi bộ - Ngày tổng hợp: {datetime.now().strftime('%d/%m/%Y')}\n")
    run_subtitle.font.name = 'Times New Roman'
    run_subtitle.font.size = Pt(12)
    run_subtitle.font.italic = True
    
    doc.add_paragraph("---------------------------\n")

    cac_mang = df["Mảng/Bộ phận"].unique()
    for mang in cac_mang:
        # Bỏ qua nếu mảng bị rỗng do lỗi dữ liệu
        if pd.isna(mang):
            continue
            
        h = doc.add_paragraph()
        run_h = h.add_run(f"I. Ý KIẾN THU THẬP TỪ MẢNG: {str(mang).upper()}")
        run_h.font.name = 'Times New Roman'
        run_h.font.size = Pt(13)
        run_h.bold = True
        
        df_sub = df[df["Mảng/Bộ phận"] == mang]
        stt = 1
        for index, row in df_sub.iterrows():
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            run_info = p.add_run(f"{stt}. Đồng chí: {row['Họ và Tên']} ({row['Chức vụ Đảng/Chính quyền']}) - [{row['Thời gian']}]\n")
            run_info.font.name = 'Times New Roman'
            run_info.font.size = Pt(11)
            run_info.bold = True
            
            run_content = p.add_run(f"   Nội dung: {row['Nội dung ý kiến']}\n")
            run_content.font.name = 'Times New Roman'
            run_content.font.size = Pt(11)
            stt += 1
        doc.add_paragraph()
    doc.save(WORD_FILE)
    return True

# --- GIAO DIỆN WEB ---
st.set_page_config(page_title="Lấy ý kiến Chi bộ", page_icon="☭", layout="centered")

st.markdown('<p style="color:red; font-size:17px;"><b>☭ HỆ THỐNG THU THẬP Ý KIẾN CHI BỘ 5</b></p>', unsafe_allow_html=True)
st.markdown('<p style="color:blue; font-size:15px;">Xin chào các đồng chí, vui lòng điền thông tin và đóng góp ý kiến xây dựng Chi bộ.</p>', unsafe_allow_html=True)

danh_sach_mang = ["Công tác Đảng", "Chuyên môn", "Đoàn trường", "Giáo dục thể chất", "Tin học", "Văn phòng", "Ý kiến khác"]
mang_selected = st.selectbox("1. Chọn Mảng/Bộ phận đóng góp ý kiến:", danh_sach_mang)

ho_ten = st.text_input("2. Họ và tên (Có thể để trống để giữ bí mật danh tính):", key="input_hoten")
chuc_vu = st.text_input("3. Chức vụ (Đảng/Chính quyền - Ví dụ: Bí thư Chi..):", key="input_chucvu")
noi_dung = st.text_area("4. Nội dung ý kiến đóng góp tâm huyết:", key="input_noidung")

btn_submit = st.button("🚀 GỬI Ý KIẾN ĐỐNG GÓP", type="primary")

if btn_submit:
    if not noi_dung.strip():
        st.error("Vui lòng nhập nội dung ý kiến, không được bỏ trống!")
    else:
        # Gọi hàm lưu trực tuyến lên mây
        luu_y_kien_vao_sheets(mang_selected, ho_ten, chuc_vu, noi_dung)
        st.success("Gửi ý kiến thành công! Cảm ơn đóng góp của đồng chí.")
        st.balloons()

# --- KHÔNG GIAN BẢO MẬT DÀNH CHO CẤP ỦY ---
st.markdown("---")
st.subheader("🔑 Khu vực dành cho Chi ủy (Bảo mật)")

mat_khau_admin = st.text_input("Nhập mật khẩu Admin để tổng hợp dữ liệu:", type="password")

if mat_khau_admin == "Chibo2026":
    st.info("Mật khẩu chính xác. Đồng chí có thể xem dữ liệu và xuất báo cáo Word bên dưới.")
    
    # Đọc dữ liệu từ Google Sheets về để hiển thị và xử lý file Word
    try:
        df_view = conn.read(ttl=0)
        df_view = df_view.dropna(how='all') # Loại bỏ hàng rỗng
    except Exception:
        df_view = pd.DataFrame()
        
    if not df_view.empty:
        st.dataframe(df_view)
        
        if st.button("🔄 Bước 1: Tiến hành tổng hợp vào file tonghop.docx"):
            if xuat_file_word_tu_sheets(df_view):
                st.success("Đã biên tập và nhóm ý kiến thành công vào file Word!")
                
        if os.path.exists(WORD_FILE):
            with open(WORD_FILE, "rb") as f:
                st.download_button(
                    label="📥 Bước 2: Tải file tonghop.docx về máy tính",
                    data=f,
                    file_name="tonghop.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.warning("Hiện tại chưa có ý kiến nào được đóng góp trên hệ thống Google Sheets.")
elif mat_khau_admin != "":
    st.error("Mật khẩu không đúng! Vui lòng kiểm tra lại.")
